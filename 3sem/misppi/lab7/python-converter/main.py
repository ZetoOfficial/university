import argparse
import json
from datetime import date

import docx
from pandas import DataFrame
from pydantic import BaseModel


class Patient(BaseModel):
    fio: str
    birthdate: date
    doctor_visit_date: date
    diagnosis: str
    sex: bool


class FileInfo(BaseModel):
    path: str
    file_type: str


def convert_to_xlsx(data: Patient, path: str) -> FileInfo:
    df = DataFrame(data.dict(), index=[0])
    df.to_excel(path)
    return FileInfo(path=path, file_type='excel')


def convert_to_word(data: Patient, path: str) -> FileInfo:
    doc = docx.Document()
    doc.add_paragraph('Patient INFO')
    doc.add_paragraph(f'Ф.И.О. пациента: {data.fio}')
    doc.add_paragraph(f'Дата рождения: {data.birthdate}')
    doc.add_paragraph(f'Дата посещения врача: {data.doctor_visit_date}')
    doc.add_paragraph(f'Диагноз: {data.diagnosis}')
    doc.add_paragraph(f'Пол: {"мужчина" if data.sex else "женщина"}')
    doc.save(path)
    return FileInfo(path=path, file_type='excel')


if __name__ == '__main__':
    parser = argparse.ArgumentParser()
    parser.add_argument("--path", help="Save path")
    parser.add_argument("--data", help="Convert data")
    args = parser.parse_args()
    dict_data = json.loads(args.data.replace("'", '"'))
    data = Patient.parse_obj(dict_data)
    if args.path.endswith('.xlsx'):
        convert_to_xlsx(data, args.path)
    if args.path.endswith('.docx'):
        convert_to_word(data, args.path)
